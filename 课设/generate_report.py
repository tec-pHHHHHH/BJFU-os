#!/usr/bin/env python
"""生成操作系统A课程设计报告 .docx"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ============================================================
# 页面设置
# ============================================================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)  # 小四号
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.0  # 单倍行距

def add_paragraph_cn(doc, text, font_name='宋体', size=Pt(12), bold=False, alignment=None, font_name_en='Times New Roman'):
    """添加段落，设置中英文字体"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name_en
    run.font.size = size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    return p

def add_heading_custom(doc, text, level=1):
    """添加标题"""
    if level == 0:
        add_paragraph_cn(doc, text, font_name='黑体', size=Pt(22), bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)  # 二号
    elif level == 1:
        add_paragraph_cn(doc, text, font_name='黑体', size=Pt(16), bold=True,
                         alignment=WD_ALIGN_PARAGRAPH.CENTER)  # 三号
    elif level == 2:
        add_paragraph_cn(doc, text, font_name='黑体', size=Pt(15), bold=True)  # 小三号
    elif level == 3:
        add_paragraph_cn(doc, text, font_name='黑体', size=Pt(14), bold=True)  # 四号
    else:
        add_paragraph_cn(doc, text, font_name='黑体', size=Pt(12), bold=True)

def add_body(doc, text):
    """添加正文段落（宋体小四）"""
    return add_paragraph_cn(doc, text, font_name='宋体', size=Pt(12), bold=False)

def add_code(doc, text):
    """添加代码块（Consolas 五号）"""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(10.5)  # 五号
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.line_spacing = 1.0

# ============================================================
# 封面
# ============================================================
for _ in range(3):
    doc.add_paragraph()

add_heading_custom(doc, '北京林业大学', level=0)
doc.add_paragraph()
add_heading_custom(doc, '2025学年—2026学年第二学期  实验报告书', level=1)

for _ in range(2):
    doc.add_paragraph()

info_items = [
    ('专    业：', '数字媒体技术'),
    ('班    级：', '数媒24-1-2'),
    ('姓    名：', '_______________'),
    ('学    号：', '_______________'),
    ('实验地点：', '_______________'),
    ('任课教师：', '范新'),
    ('实验题目：', '持久化操作系统核心模拟器'),
    ('实验环境：', 'C语言（C99）+ GCC + MSYS2 + Windows 11'),
]

for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}    {value}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)  # 四号
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.5

doc.add_page_break()

# ============================================================
# 一、实验目的
# ============================================================
add_heading_custom(doc, '一、实验目的', level=2)

add_body(doc, '本课程设计旨在通过实现一个持久化操作系统核心模拟器，深入理解操作系统的核心概念和实现机制。具体目标如下：')

purposes = [
    '1. 掌握进程控制块（PCB）的设计与实现，理解进程的创建、撤销、阻塞、唤醒、挂起、恢复等状态转换；',
    '2. 掌握多级反馈队列（MLFQ）调度算法，理解优先级调度、时间片轮转、降级机制；',
    '3. 掌握动态分区内存管理，理解首次适应（FF）、最佳适应（BF）、最坏适应（WF）三种分配算法，以及碎片紧缩（compact）机制；',
    '4. 掌握状态持久化技术，使用二进制格式保存和恢复系统完整状态；',
    '5. 掌握多线程并发编程，理解生产者-消费者模型、互斥锁（mutex）、条件变量（cond）及文件锁机制；',
    '6. 掌握操作系统可视化展示方法，以ASCII图形化方式呈现进程树、内存布局和调度队列。',
]

for purpose in purposes:
    add_body(doc, purpose)

doc.add_paragraph()

# ============================================================
# 二、实验内容
# ============================================================
add_heading_custom(doc, '二、实验内容', level=2)

add_body(doc, '设计并实现一个模拟操作系统核心功能的交互式程序，重点覆盖进程管理、内存管理、调度器、系统状态持久化、多线程并发控制和可视化展示。具体包含以下七大模块：')

modules_overview = [
    ('模块一：用户管理（8分）', '用户注册、登录、登出，密码哈希存储，3次错误登录锁定及锁定状态持久化。'),
    ('模块二：进程管理（20分）', '10个进程操作命令：create_pcb、kill_pcb、block_pcb、wakeup_pcb、show_pcb、list_pcb、ptree、suspend、resume、renice。init进程（PID=1）自动创建，kill递归清理子进程并回收内存。'),
    ('模块三：调度器（8分）', '三级MLFQ：Q0(prio 0-3, q=2)、Q1(prio 4-7, q=4)、Q2(prio 8-15, q=8)。step展示4步完整决策链路，start_sched启动自动调度线程，stop_sched/restart_sched控制启停。'),
    ('模块四：内存管理（16分）', '8个内存命令：alloc、free_mem、show_mem、compact、mem_stat、set_alloc_algo、pgfault、swap_out。支持FF/BF/WF三种分配算法，碎片率统计（1-max_contiguous/total_free），自动合并相邻空闲块。'),
    ('模块五：持久化（20分）', '二进制格式（OS_MAGIC校验）保存/恢复进程、MLFQ队列、内存分配、系统时钟的完整状态。启动时自动检测并载入，冷启动正常初始化。'),
    ('模块六：多线程并发（18分）', '四线程架构：前端线程（生产者，读stdin入队）、后端线程（消费者，出队执行）、调度线程（自动tick）、守护线程（检测外部文件变化自动同步）。消息队列使用mutex+cond实现线程安全。文件锁（O_CREAT|O_EXCL）支持多实例并发。'),
    ('模块七：可视化（10分）', 'overview命令展示三大板块：进程树（含PID、状态、优先级、内存）、内存ASCII分布图、MLFQ三级队列状态（含空队列显示）。'),
]

for title, desc in modules_overview:
    p = doc.add_paragraph()
    run_title = p.add_run(title)
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(12)
    run_title.bold = True
    run_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run_desc = p.add_run(desc)
    run_desc.font.name = 'Times New Roman'
    run_desc.font.size = Pt(12)
    run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.paragraph_format.line_spacing = 1.0

doc.add_paragraph()

# ============================================================
# 三、实验环境
# ============================================================
add_heading_custom(doc, '三、实验环境', level=2)

env_items = [
    '• 操作系统：Windows 11 Home China',
    '• 编译器：GCC 15.2.0（MSYS2/MinGW-w64）',
    '• 语言标准：C99（ISO/IEC 9899:1999）',
    '• 构建工具：GNU Make',
    '• 编辑器：Visual Studio Code',
    '• 第三方库：pthread（POSIX线程库）',
    '• 版本控制：无（本地开发）',
]

for item in env_items:
    add_body(doc, item)

doc.add_paragraph()

# ============================================================
# 四、实现方法
# ============================================================
add_heading_custom(doc, '四、实现方法', level=2)

add_body(doc, '本项目采用模块化设计，共8个源文件和6个头文件，各模块职责分明。系统启动流程为：初始化系统状态 → 加载用户数据库 → 自动加载持久化文件 → 启动后端消费者线程 → 主线程运行前端生产者循环。')

doc.add_paragraph()

# 4.1 整体架构
add_heading_custom(doc, '4.1 整体架构设计', level=3)

add_body(doc, '系统采用"前端-后端"分离的多线程架构：')

arch_text = '''前端线程（frontend_thread，运行于main线程）：
  负责阻塞读取用户输入（stdin），将原始命令字符串封装为Command结构体，
  通过mq_put()推入线程安全的消息队列（MessageQueue），然后立即返回继续读取下一条输入。
  该线程不处理任何业务逻辑。

后端线程（backend_thread，独立pthread）：
  通过mq_get()从消息队列阻塞取出命令，解析命令名称和参数，
  分发到对应模块（进程管理/调度器/内存管理/持久化/可视化）执行。
  命令执行完毕后检查autosave状态，决定是否自动保存。

调度线程（scheduler_thread_func，独立pthread）：
  由start_sched命令启动，每1秒自动执行一次scheduler_tick()。
  通过sched_paused标志 + pthread_cond_wait/pthread_cond_signal实现暂停/恢复控制。

守护线程（reload_daemon_thread，独立pthread）：
  登录成功后由concurrency_start_daemon()启动，每2秒检查os_state.bin的mtime。
  若检测到外部修改则自动调用persistence_load()同步状态。'''

add_body(doc, arch_text)

doc.add_paragraph()

# 4.2 数据结构
add_heading_custom(doc, '4.2 核心数据结构', level=3)

add_body(doc, '（1）PCB（进程控制块）— os_types.h:53-84')

pcb_text = '''PCB结构体包含以下字段：
  • pid / ppid：进程ID和父进程ID
  • name[MAX_NAME_LEN] / owner[MAX_NAME_LEN]：进程名称和所属用户
  • state：进程状态枚举（RUNNING / READY / BLOCKED / SUSPENDED）
  • priority：优先级（0最高 ~ 15最低）
  • mem_start / mem_size：内存分配信息（KB）
  • time_slice / total_ticks / queue_level：调度信息
  • parent / first_child / next_sibling：进程树三指针（父、长子、兄弟）
  • next：全局进程链表指针
  • queue_next / queue_prev：MLFQ双向循环链表指针'''
add_body(doc, pcb_text)

add_body(doc, '（2）MessageQueue（消息队列）— os_types.h:166-173')

mq_text = '''消息队列基于单向链表实现，包含以下同步原语：
  • pthread_mutex_t mutex：保护队列所有操作的互斥锁
  • pthread_cond_t cond：条件变量，用于消费者在队列空时阻塞等待
  • int quit_flag：退出标志，使消费者在收到退出信号后正常结束'''
add_body(doc, mq_text)

add_body(doc, '（3）MLFQ_Level（调度队列）— os_types.h:119-127')

mlfq_text = '''每级调度队列为双向循环链表：
  • PCB *head / *tail：队头队尾指针（head是下一个要调度的进程）
  • int count：队列中进程数
  • int quantum：该级时间片（Q0=2, Q1=4, Q2=8 ticks）
  • int prio_min / prio_max：该级优先级范围'''
add_body(doc, mlfq_text)

doc.add_paragraph()

# 4.3 模块一：用户管理
add_heading_custom(doc, '4.3 模块一：用户管理', level=3)

add_body(doc, '代码位置：src/user.c（约280行）')

add_body(doc, '（1）注册（user_register）：验证用户名合法性（仅允许字母、数字、下划线）和密码非空，检查用户名是否重复，将用户信息存入全局链表并写入data/users.db文件。')

add_body(doc, '（2）密码哈希（simple_hash）：采用自定义ARX哈希算法（Add-Rotate-XOR），将明文密码转换为64字符十六进制字符串存储，不可逆。验证登录时对输入密码哈希后比对，原始密码不存储。')

add_body(doc, '（3）登录锁定：每个用户维护login_attempts计数器，连错3次后is_locked置1，即使正确密码也无法登录。锁定状态持久化到users.db，重启后仍有效。')

add_body(doc, '（4）登录（user_login）：校验用户名是否存在、账户是否锁定、密码哈希是否匹配。成功后设置current_user，启动守护线程，重置错误计数。')

add_body(doc, '（5）登出（user_logout）：清空current_user，停止守护线程。')

doc.add_paragraph()

# 4.4 模块二：进程管理
add_heading_custom(doc, '4.4 模块二：进程管理', level=3)

add_body(doc, '代码位置：src/process.c（约511行）')

add_body(doc, '（1）create_pcb：首次调用时自动创建init进程（PID=1, prio=0, mem=64KB）。之后创建的进程owner为当前登录用户，默认父进程为当前运行进程或init。创建后调用pcb_enqueue_ready()根据优先级入队对应MLFQ队列。')

add_body(doc, '（2）kill_pcb：采用递归策略——先递归kill所有子进程，再从MLFQ队列中出队（若就绪态），释放其已分配内存（从alloc_list移除并归还free_list），从父进程的first_child链表和全局process_list中移除，最后释放PCB内存。init进程不允许被kill。')

add_body(doc, '（3）block_pcb / wakeup_pcb：阻塞态从MLFQ出队，唤醒态重新入队。各状态转换有合法性检查（如不能阻塞已阻塞进程、不能唤醒非阻塞进程）。')

add_body(doc, '（4）suspend / resume：挂起态从MLFQ移除，恢复态重新入队。与阻塞的区别是语义层面（挂起是主动从调度移除）。')

add_body(doc, '（5）renice：修改优先级后若进程在就绪队列，需先出队再按新优先级重新入队（可能跨队列）。')

add_body(doc, '（6）pcb_enqueue_ready：按priority_to_queue()映射到Q0/Q1/Q2，双向循环链表尾插，设置time_slice为该队列quantum。')

add_body(doc, '（7）pcb_ptree：从init开始递归遍历进程树，使用"+- "、"|  "等ASCII字符绘制树形结构，显示每个进程的状态标记（RDY/BLK/SUS/RUN）、优先级和内存占用。')

doc.add_paragraph()

# 4.5 模块三：调度器
add_heading_custom(doc, '4.5 模块三：调度器', level=3)

add_body(doc, '代码位置：src/scheduler.c（约306行）')

add_body(doc, '（1）核心调度 tick（scheduler_tick）：从Q0开始扫描第一个非空队列 → 取其队头出队 → 执行min(remaining_time_slice, queue_quantum)个tick → 更新system_clock和进程total_ticks → 若时间片耗尽则降级到下一级队列（Q0→Q1、Q1→Q2，Q2到底不再降），给满新队列的时间片；若未耗尽则回原队列队尾（RR轮转）。')

add_body(doc, '（2）step单步执行：展示完整4步决策链——Step1当前队列快照 → Step2选程逻辑（从Q0扫描，为什么选这个而不选那个） → Step3执行（调用scheduler_tick） → Step4最终状态。')

add_body(doc, '（3）自动调度线程：scheduler_start创建调度线程 → 循环执行tick+usleep(1000ms) → scheduler_stop通过pthread_cond_wait阻塞 → scheduler_restart通过pthread_cond_signal唤醒。')

doc.add_paragraph()

# 4.6 模块四：内存管理
add_heading_custom(doc, '4.6 模块四：内存管理', level=3)

add_body(doc, '代码位置：src/memory.c（约467行）')

add_body(doc, '（1）alloc：调用mem_find_free()按当前算法（FF/BF/WF）查找合适空闲块 → 创建AllocBlock → 从FreeBlock中切割（刚好用完则删除空闲块，否则修改start和size） → 更新进程mem_start/mem_size。')

add_body(doc, '（2）三种分配算法（mem_find_free）：First-Fit遍历返回第一个size≥要求的块（最快）；Best-Fit遍历所有，返回最小满足块（最省空间）；Worst-Fit遍历所有，返回最大满足块（剩余部分仍可用）。')

add_body(doc, '（3）free_mem + 自动合并：在alloc_list中查找对应块 → 移除 → 创建FreeBlock → 调用free_list_insert_sorted()按地址排序插入 → 遍历相邻块，若curr->start + curr->size == next->start则合并为一个大块。')

add_body(doc, '（4）compact碎片紧缩：将所有已分配块按start排序 → 从地址0开始逐个紧排（第i块start = 前i-1块size总和） → 更新对应PCB的mem_start → 清空旧free_list，所有剩余空间合并为一块大FreeBlock。')

add_body(doc, '（5）mem_stat碎片率：遍历free_list统计total_free和max_contiguous_free → 碎片率 = 1 - max_free/total_free → 若碎片率>5%且空闲块数>1则发出警告。')

add_body(doc, '（6）swap_out：将指定进程的AllocBlock移除并归还free_list，进程mem_start=-1/mem_size=0标记为已换出（模拟写入磁盘交换空间）。')

doc.add_paragraph()

# 4.7 模块五：持久化
add_heading_custom(doc, '4.7 模块五：持久化', level=3)

add_body(doc, '代码位置：src/persistence.c（约364行）')

add_body(doc, '（1）二进制文件格式（逐段fwrite）：Header（9个int：magic=0x4F533230, version, pcb_count, next_pid, system_clock, sched_running, alloc_algo, free_count, alloc_count）→ PCB Records（固定长度PCBRecord结构体数组）→ MLFQ State（每级队列count + PID序列）→ Free Blocks（start, size）→ Alloc Blocks（start, size, pid）。')

add_body(doc, '（2）persistence_save：统计各段计数 → fopen("wb")二进制写 → 依次写入5个段 → fclose → 置dirty=0。')

add_body(doc, '（3）persistence_load：fopen("rb")二进制读 → 读取Header并校验magic → clear_state()清空当前状态 → 读PCB Records并重建进程树（parent/child/sibling三指针） → 读MLFQ并重建双向循环链表 → 读Free/Alloc Blocks重建链表 → 恢复system_clock等字段。')

add_body(doc, '（4）冷启动处理：若fopen失败或magic不匹配 → 打印warning → 返回错误码 → 系统以init_system()的默认状态启动。')

doc.add_paragraph()

# 4.8 模块六：多线程并发
add_heading_custom(doc, '4.8 模块六：多线程并发', level=3)

add_body(doc, '代码位置：src/main.c（消息队列+前后端线程）、src/scheduler.c（调度线程）、src/concurrency.c（守护线程+文件锁+autosave）')

add_body(doc, '（1）生产者-消费者模型（main.c:25-78, 169-199, 205-587）：')

add_body(doc, '  消息队列：基于单向链表，mq_put()加锁 → 节点尾插 → cond_signal唤醒消费者 → 解锁；mq_get()加锁 → 若队列空且未quit则cond_wait阻塞 → 取出队头节点 → 解锁 → 返回命令。')

add_body(doc, '  前端（生产者）：fgets阻塞读stdin → mq_put()入队（不管命令内容，直接返回继续读）。')

add_body(doc, '  后端（消费者）：mq_get()阻塞取命令 → 解析 → if-else分发到各模块函数。')

add_body(doc, '（2）调度线程（scheduler.c:141-165）：独立pthread，while(s->sched_running)循环，每次tick后usleep(1000ms)。sched_paused时pthread_cond_wait阻塞，restart时pthread_cond_signal唤醒。')

add_body(doc, '（3）守护线程（concurrency.c:77-122）：登录成功自动启动，每2秒stat()检查os_state.bin的mtime，变化时若调度器运行中则先暂停→persistence_load()覆盖内存状态→打印同步信息→恢复调度器。退出时concurrency_stop_daemon()设reload_running=0+pthread_join等待结束。')

add_body(doc, '（4）文件锁（concurrency.c:27-51）：O_CREAT|O_EXCL原子操作——成功创建锁文件=获取锁，EEXIST=锁被占用→usleep(200ms)重试最多10次。concurrency_unlock()删除锁文件。用于autosave和手动save时保护二进制文件不被并发写入破坏。')

doc.add_paragraph()

# 4.9 模块七：可视化
add_heading_custom(doc, '4.9 模块七：可视化', level=3)

add_body(doc, '代码位置：src/visualization.c（约180行）')

add_body(doc, 'overview命令展示三个板块：')

add_body(doc, '板块一（进程树）：从init开始递归调用tree_print()，以ASCII字符"+- "和"|  "绘制树形缩进结构，每个节点显示name(PID) [状态, prio=优先级, mem=内存KB]。孤儿进程（父进程不存在）单独列出。')

add_body(doc, '板块二（内存ASCII分布图）：将free_list和alloc_list的块按start排序后，逐段绘制：已分配块显示##name(size)##，空闲块显示--free(size)--，间隙显示--gap(size)--。')

add_body(doc, '板块三（MLFQ队列）：遍历三级队列，显示每级prio范围、quantum值，队列中进程以"name(PID=#,rem=剩余时间片)"格式按序打印（空队列显示"(empty)"）。')

doc.add_paragraph()

# ============================================================
# 五、实验结果及分析
# ============================================================
add_heading_custom(doc, '五、实验结果及分析', level=2)

# 5.1
add_heading_custom(doc, '5.1 模块一：用户管理测试', level=3)

add_body(doc, '测试命令及结果：')

test1_input = """# 正常注册
register admin 123456
  => [Register] User 'admin' created.

# 重复用户名
register admin xyz
  => [Register] FAIL: user 'admin' already exists.

# 非法字符
register admin@test 123
  => [Register] FAIL: invalid username (alphanumeric + underscore only).

# 正常登录
login admin 123456
  => [System] Welcome 'admin'! Login successful.

# 错误密码
login admin wrong
  => [Login] FAIL: wrong password.

# 3次锁定
login admin wrong1  →  FAIL: wrong password.
login admin wrong2  →  FAIL: wrong password.
login admin wrong3  →  [System] Account 'admin' is now LOCKED!
login admin 123456  →  [Login] FAIL: account 'admin' is locked."""

add_code(doc, test1_input)

add_body(doc, '结果分析：用户注册/登录/锁定功能均正常。密码以哈希形式存储在users.db中，原始密码不可见。锁定状态在重启后保持，验证了持久化机制。')

doc.add_paragraph()

# 5.2
add_heading_custom(doc, '5.2 模块二：进程管理测试', level=3)

add_body(doc, '测试命令及结果：')

test2_input = """create_pcb jobA 3
  => [Create] init (PID=1, prio=0, mem=64KB) auto-created.
  => [Create] jobA (PID=2, prio=3) -- parent PID=1

create_pcb jobB 8
  => [Create] jobB (PID=3, prio=8) -- parent PID=1

list_pcb
  => PID  NAME   STATE   PRI  MEM   TICKS  OWNER
     1  init   READY    0    64K   0     system
     2  jobA   READY    3     0K   0     admin
     3  jobB   READY    8     0K   0     admin

block_pcb 2 → show_pcb 2 → State: BLOCKED
wakeup_pcb 2 → show_pcb 2 → State: READY
suspend 3 → show_pcb 3 → State: SUSPENDED
resume 3 → show_pcb 3 → State: READY
renice 2 10 → [Renice] priority 3 -> 10 (Q0->Q2)"""

add_code(doc, test2_input)

add_body(doc, '结果分析：10个进程管理命令均正确执行。init自动创建、进程状态转换（READY↔BLOCKED↔SUSPENDED）、优先级修改及跨队列迁移、权限检查等功能均通过验证。')

doc.add_paragraph()

# 5.3
add_heading_custom(doc, '5.3 模块三：调度器测试', level=3)

add_body(doc, '测试命令及结果：')

test3_input = """step
=> +==================================+
   |       STEP Single-Step           |
   +==================================+
   [Step 1] Current queue snapshot:
   Q0(prio 0-3, quantum=2): init(PID=1,rem=2)
   Q1: (empty)
   Q2: jobA(PID=2,rem=8) -> jobB(PID=3,rem=8)
   [Step 2] Selection logic:
   -> Scan Q0: non-empty (1 processes)
   -> Selected Q0 head: init(PID=1, ...)
   [Step 3] Execution:
   [Sched tick=2] Selected: init(PID=1) from Q0
   >>> Timeslice exhausted! Demotion: Q0 -> Q1 (new timeslice=4)
   [Step 4] Final state:"""

add_code(doc, test3_input)

add_body(doc, '结果分析：step展示了完整的4步决策过程（队列快照→选程逻辑→执行→最终状态）。时间片耗尽后init从Q0降级到Q1，验证了MLFQ的核心降级机制。start_sched/stop_sched/restart_sched启停控制正常。')

doc.add_paragraph()

# 5.4
add_heading_custom(doc, '5.4 模块四：内存管理测试', level=3)

add_body(doc, '测试命令及结果：')

test4_input = """alloc 200 → [Alloc] Success! addr=0 KB, size=200 KB (algo: First-Fit)
alloc 100 → [Alloc] Success! addr=200 KB, size=100 KB (algo: First-Fit)
set_alloc_algo 1 → [Algo] Switched to: Best-Fit (BF)
alloc 50  → [Alloc] Success! addr=300 KB, size=50 KB (algo: Best-Fit)
free_mem 200 → [Free] addr=200 KB, size=100 KB returned to free list.
compact → [Compact] Done! All allocated blocks moved to low addresses.
mem_stat → Fragmentation: 0.0% (compact后碎片率为0)"""

add_code(doc, test4_input)

add_body(doc, '结果分析：alloc/free_mem正常工作，三种分配算法切换有效。compact碎片紧缩后所有已分配块紧密排列在低地址，碎片率降为0%。碎片率计算公式1-max_contiguous/total_free正确实现。')

doc.add_paragraph()

# 5.5
add_heading_custom(doc, '5.5 模块五：持久化测试', level=3)

add_body(doc, '测试命令及结果：')

test5_input = """# 创建测试环境
create_pcb worker1 3
create_pcb worker2 8
alloc 200
save → [Save] State saved to data/os_state.bin
       Processes: 3 | Free blocks: 1 | Alloc blocks: 2 | Clock: 0
exit → [System] Shutting down...

# 重新启动程序
login admin 123456
list_pcb → 三个进程完整恢复（init, worker1, worker2）
show_mem → 内存分配状态恢复
overview → MLFQ队列、进程树、内存布局全部正确恢复"""

add_code(doc, test5_input)

add_body(doc, '结果分析：二进制持久化save/load功能正常。重启后进程数量、名称、PID、父子关系、MLFQ队列顺序、内存分配状态、系统时钟完整恢复。冷启动（无os_state.bin）时正常初始化，不崩溃。文件经file命令验证为二进制而非ASCII文本。')

doc.add_paragraph()

# 5.6
add_heading_custom(doc, '5.6 模块六：多线程并发测试', level=3)

add_body(doc, '测试方法：两个终端同时运行os_sim.exe。')

add_body(doc, '终端A：login → create_pcb from_A 5 → alloc 200 → save。终端B：login → 2秒后自动输出"[Daemon] External update detected, state auto-synced." → list_pcb可见from_A进程。')

add_body(doc, '并发压力测试：在start_sched日志持续刷屏时输入命令（list_pcb/overview/step），命令正常响应不崩溃，验证前后台线程分离和消息队列互斥锁保护有效。')

add_body(doc, '结果分析：四个线程（前端生产者、后端消费者、调度、守护）协同工作正常。文件锁机制保证了多实例写入安全，守护线程实现了跨终端的自动状态同步。')

doc.add_paragraph()

# 5.7
add_heading_custom(doc, '5.7 模块七：可视化测试', level=3)

add_body(doc, '测试命令及结果（见下方完整输出示例）：')

add_body(doc, 'overview命令输出包含三个清晰板块，使用纯ASCII字符绘制。操作后（step/create/kill/compact）再次overview，三个板块均正确反映最新状态。')

doc.add_paragraph()

# ============================================================
# 六、结论分析
# ============================================================
add_heading_custom(doc, '六、结论分析', level=2)

add_body(doc, '本课程设计成功实现了一个功能完整的持久化操作系统核心模拟器，涵盖PPT要求的全部7个模块（100分制），核心成果如下：')

conclusions = [
    '1. 进程管理：实现了完整的10条进程管理命令和4种状态（READY/RUNNING/BLOCKED/SUSPENDED）之间的转换，进程树（parent/first_child/next_sibling）和全局链表双索引结构保证了操作的完整性。kill递归清理子进程和回收内存，保证了资源不泄漏。',
    '2. MLFQ调度：三级队列（Q0交互/Q1普通/Q2批处理）正确实现了"高优先级优先、时间片耗尽降级、同队列RR轮转"的调度策略。step命令的4步决策展示使调度过程透明可验证。',
    '3. 内存管理：FF/BF/WF三种分配算法的切换机制灵活，auto-merge相邻空闲块和compact碎版紧缩有效降低了外部碎片。碎片率公式符合PPT要求。',
    '4. 持久化：自定义二进制文件格式（magic+version+分段结构）保证了数据完整性和格式版本兼容。启动自动加载+手动load+守护自动同步三条路径共享同一核心函数，代码复用率高。',
    '5. 多线程并发：生产者-消费者模型成功分离了输入处理和执行逻辑，mutex+cond保证了线程安全。O_CREAT|O_EXCL文件锁实现了轻量级进程间互斥，无需引入数据库或Socket。',
    '6. 可视化：纯ASCII字符输出保证了Windows终端的兼容性，进程树/内存图/MLFQ队列三大板块信息完整，满足PPT"缺一不可"的硬性要求。',
    '',
    '存在的问题与改进方向：',
    '• 当前调度器不支持优先级老化（boost），长时间在低优先级队列的进程可能饥饿；',
    '• 内存分配未实现虚拟内存的完整页表机制（仅swap_out模拟）；',
    '• 文件锁基于轮询重试而非futex等内核级锁，高并发下效率较低；',
    '• 可扩展为CS架构，通过Socket实现真正的远程多用户并发访问。',
    '',
    '通过本次课程设计，深刻理解了操作系统的四大核心功能（进程管理、内存管理、文件系统接口、并发控制）在实践中的具体实现方式，掌握了C语言多线程编程（pthread）、二进制数据序列化、动态内存管理、模块化软件架构等工程技能。',
]

for conclusion in conclusions:
    add_body(doc, conclusion)

doc.add_paragraph()

# ============================================================
# 附录（可选）
# ============================================================
add_heading_custom(doc, '七、附录：项目文件结构', level=2)

file_structure = """课设/
├── include/
│   ├── os_types.h          # 全局数据结构定义（PCB/MLFQ/MessageQueue/SystemState等）
│   ├── user.h              # 用户管理模块接口
│   ├── process.h           # 进程管理模块接口
│   ├── scheduler.h         # 调度器模块接口
│   ├── memory.h            # 内存管理模块接口
│   ├── persistence.h       # 持久化模块接口
│   ├── concurrency.h       # 并发模块接口
│   └── visualization.h     # 可视化模块接口
├── src/
│   ├── main.c              # 系统入口 + 消息队列 + 前端/后端线程 + 命令分发
│   ├── user.c              # 用户注册/登录/登出/锁定/密码哈希
│   ├── process.c           # 10条进程管理命令 + MLFQ入队/出队
│   ├── scheduler.c         # MLFQ调度核心tick + step + auto线程
│   ├── memory.c            # alloc/free/compact/mem_stat + FF/BF/WF
│   ├── persistence.c       # save/load 二进制序列化/反序列化
│   ├── concurrency.c       # 守护线程 + 文件锁 + autosave
│   └── visualization.c     # overview三板块ASCII展示
├── data/                   # 运行时数据（users.db, os_state.bin, os_state.lock）
├── build/                  # 编译输出
├── Makefile                # 构建脚本
└── 测试流程.md              # 完整手工测试指南"""

add_code(doc, file_structure)

doc.add_paragraph()
add_body(doc, '注：完整源代码约3200行C语言代码，遵循C99标准。所有用户可见字符串使用英文，无需UTF-8终端支持。')

# ============================================================
# 保存
# ============================================================
output_path = r'd:\file\study\big_2_xia\O_System\课设\操作系统A课程设计报告.docx'
doc.save(output_path)
print(f'报告已生成: {output_path}')
print('Done!')
